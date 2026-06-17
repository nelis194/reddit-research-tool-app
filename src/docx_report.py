"""Bouw een net opgemaakt Word-rapport (.docx) uit een ``AnalysisResult``.

Bevat alle hoofdstukken: samenvatting, pijnpunten, frustraties, gewenste
uitkomsten, geslaagde/mislukte oplossingen, koopbezwaren/-motieven, merken,
concurrenten, voice-of-customer, persona's, taalgebruik, content angles, ad
hooks, FAQ's, before/after en (indien aanwezig) de AI-inzichten van Claude.
"""

from __future__ import annotations

from io import BytesIO
from typing import Dict, List

from docx import Document
from docx.shared import Pt, RGBColor

from .analyzer import AnalysisResult
from .utils import now_iso

BLUE = RGBColor(0x1E, 0x40, 0xAF)
INK = RGBColor(0x0F, 0x17, 0x2A)
MUTED = RGBColor(0x47, 0x55, 0x69)


def _color_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BLUE
    return h


def _theme_section(doc: Document, title: str, items: List[Dict]) -> None:
    _color_heading(doc, title, 1)
    if not items:
        p = doc.add_paragraph("Geen resultaten gevonden voor dit onderwerp.")
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = MUTED
        return
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(str(it.get("theme", "")).strip())
        run.bold = True
        count = it.get("count")
        if count:
            c = p.add_run(f"  ({count}x)")
            c.font.color.rgb = MUTED
        example = it.get("example")
        if example:
            q = doc.add_paragraph(style="List Bullet 2")
            qr = q.add_run(f"“{example}”")
            qr.italic = True
            qr.font.color.rgb = MUTED


def _bullets(doc: Document, title: str, items: List[str]) -> None:
    if not items:
        return
    _color_heading(doc, title, 1)
    for it in items:
        doc.add_paragraph(str(it), style="List Bullet")


def build_docx(
    result: AnalysisResult,
    posts_count: int = 0,
    comments_count: int = 0,
) -> bytes:
    """Genereer het Word-rapport en geef de bytes terug."""
    doc = Document()

    # Basis-lettertype.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ---------------------------------------------------------------- titelpagina
    title = doc.add_heading("Reddit Research — Onderzoeksrapport", level=0)
    for run in title.runs:
        run.font.color.rgb = BLUE

    kw = ", ".join(result.keyword_context) or "(geen opgegeven)"
    sub = doc.add_paragraph()
    r = sub.add_run(f"Onderwerp / keywords: {kw}")
    r.bold = True
    r.font.size = Pt(12)

    meta = doc.add_paragraph()
    meta_line = (
        f"Gegenereerd: {now_iso()[:19].replace('T', ' ')}  ·  "
        f"Posts: {posts_count}  ·  Comments: {comments_count}  ·  "
        f"Geanalyseerde documenten: {result.document_count}"
    )
    mr = meta.add_run(meta_line)
    mr.font.size = Pt(9)
    mr.font.color.rgb = MUTED

    if result.sentiment:
        s = doc.add_paragraph()
        sr = s.add_run(
            f"Algemeen sentiment: {result.sentiment.get('label', '-')} "
            f"(polariteit {result.sentiment.get('polarity', 0)})"
        )
        sr.font.color.rgb = MUTED
        sr.font.size = Pt(9)

    if result.is_medical and result.disclaimer:
        d = doc.add_paragraph()
        dr = d.add_run(result.disclaimer)
        dr.italic = True
        dr.font.size = Pt(9)

    doc.add_paragraph()

    # --------------------------------------------------------------- samenvatting
    _color_heading(doc, "1. Samenvatting", 1)
    summary = doc.add_paragraph(
        f"Dit rapport vat de Reddit-discussies samen rondom: {kw}. "
        f"Op basis van {result.document_count} geanalyseerde fragmenten zijn de "
        f"belangrijkste pijnpunten, frustraties, oplossingen, koopmotieven en "
        f"voice-of-customer-citaten uitgewerkt in de hoofdstukken hieronder."
    )
    summary  # noqa: B018

    # ----------------------------------------------------------------- hoofdstukken
    _theme_section(doc, "2. Top pijnpunten", result.top_pain_points)
    _theme_section(doc, "3. Top frustraties", result.top_frustrations)
    _theme_section(doc, "4. Gewenste uitkomsten", result.top_desired_outcomes)
    _theme_section(doc, "5. Geslaagde oplossingen", result.top_successful_solutions)
    _theme_section(doc, "6. Mislukte oplossingen", result.top_failed_solutions)
    _theme_section(doc, "7. Koopbezwaren", result.buying_objections)
    _theme_section(doc, "8. Koopmotieven", result.buying_motivations)

    # Merken & concurrenten.
    _color_heading(doc, "9. Genoemde merken / producten", 1)
    if result.mentioned_brands:
        for b in result.mentioned_brands[:25]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(str(b.get("term", ""))).bold = True
            p.add_run(f"  ({b.get('count', 0)}x)").font.color.rgb = MUTED
    else:
        doc.add_paragraph("Geen merken gedetecteerd.").runs[0].italic = True

    if result.mentioned_competitors:
        _color_heading(doc, "10. Genoemde concurrenten", 1)
        for c in result.mentioned_competitors:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(str(c.get("term", ""))).bold = True
            p.add_run(f"  ({c.get('count', 0)}x)").font.color.rgb = MUTED

    # Voice of customer.
    _color_heading(doc, "11. Voice of Customer (citaten)", 1)
    any_voc = False
    for cat, quotes in (result.voice_of_customer or {}).items():
        if not quotes:
            continue
        any_voc = True
        _color_heading(doc, cat.replace("_", " ").title(), 2)
        for q in quotes:
            p = doc.add_paragraph(style="List Bullet")
            qr = p.add_run(f"“{q}”")
            qr.italic = True
    if not any_voc:
        doc.add_paragraph("Geen representatieve citaten gevonden.").runs[0].italic = True

    # Persona's.
    _color_heading(doc, "12. Persona-clusters", 1)
    doc.add_paragraph(
        "Afgeleid uit taalgebruik en inhoud van de discussies — niet uit "
        "gebruikersaccounts."
    ).runs[0].font.color.rgb = MUTED
    if result.persona_clusters:
        for p_ in result.persona_clusters:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(str(p_.get("persona", "")).title()).bold = True
            p.add_run(f"  ({p_.get('signal_count', 0)} signalen)").font.color.rgb = MUTED
            if p_.get("example"):
                e = doc.add_paragraph(style="List Bullet 2")
                er = e.add_run(f"“{p_['example']}”")
                er.italic = True
                er.font.color.rgb = MUTED
    else:
        doc.add_paragraph("Geen duidelijke persona-signalen.").runs[0].italic = True

    # Taalgebruik.
    _color_heading(doc, "13. Taalgebruik", 1)
    if result.common_words:
        doc.add_paragraph("Veelgebruikte woorden:").runs[0].bold = True
        doc.add_paragraph(
            ", ".join(f"{w['word']} ({w['count']})" for w in result.common_words[:30])
        )
    if result.common_phrases:
        doc.add_paragraph("Veelgebruikte zinnen:").runs[0].bold = True
        for ph in result.common_phrases[:20]:
            doc.add_paragraph(
                f"“{ph['phrase']}” ({ph['count']}x)", style="List Bullet"
            )

    # Marketing-output.
    _bullets(doc, "14. Content angles", result.content_angles)
    _bullets(doc, "15. Ad hooks", result.ad_hooks)

    if result.faqs:
        _color_heading(doc, "16. FAQ's (uit echte vragen)", 1)
        for f in result.faqs:
            doc.add_paragraph(str(f.get("question", "")), style="List Bullet")

    if result.before_after:
        _color_heading(doc, "17. Before / after", 1)
        for b in result.before_after:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"“{b.get('quote', '')}”").italic = True

    # AI-inzichten (Claude), indien aanwezig.
    llm = result.llm_insights or {}
    if llm:
        _color_heading(doc, "18. AI-inzichten (Claude)", 1)
        for key, heading in [
            ("thematic_summaries", "Thematische samenvattingen"),
            ("marketing_insights", "Marketing insights"),
            ("customer_insights", "Customer insights"),
            ("content_ideas", "Content-ideeën"),
            ("landing_page_hooks", "Landing page hooks"),
            ("offer_ideas", "Offer-ideeën"),
            ("persona_clusters", "Persona-clusters (AI)"),
        ]:
            items = llm.get(key) or []
            if items:
                _color_heading(doc, heading, 2)
                for it in items:
                    doc.add_paragraph(str(it), style="List Bullet")

        cards = llm.get("persona_cards") or []
        if cards:
            _color_heading(doc, "Persona cards", 2)
            for card in cards:
                p = doc.add_paragraph()
                p.add_run(str(card.get("name", "Persona"))).bold = True
                if card.get("description"):
                    doc.add_paragraph(str(card["description"]))
                for ck, ctitle in [
                    ("pains", "Pijnpunten"),
                    ("desires", "Verlangens"),
                    ("objections", "Bezwaren"),
                    ("language", "Taalgebruik"),
                ]:
                    vals = card.get(ck) or []
                    if vals:
                        para = doc.add_paragraph(style="List Bullet")
                        para.add_run(f"{ctitle}: ").bold = True
                        para.add_run("; ".join(map(str, vals)))

        faqs = llm.get("faqs") or []
        if faqs:
            _color_heading(doc, "FAQ's met antwoord", 2)
            for f in faqs:
                if not isinstance(f, dict):
                    continue
                q = doc.add_paragraph()
                q.add_run(str(f.get("question", ""))).bold = True
                if f.get("answer"):
                    doc.add_paragraph(str(f["answer"]))

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
